from flask import Flask, render_template, request, jsonify, send_file
import os
import openai
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from datetime import datetime

app = Flask(__name__)

# Configure OpenAI for CrimsonAI hub using stable v0.28.1 API
def configure_openai():
    api_key = os.getenv('AZURE_OPENAI_API_KEY')
    endpoint = os.getenv('AZURE_OPENAI_ENDPOINT')
    
    if not api_key or not endpoint:
        raise ValueError("Azure OpenAI credentials not configured")
    
    openai.api_type = "azure"
    openai.api_key = api_key
    openai.api_base = endpoint
    openai.api_version = "2024-12-01-preview"

@app.route('/')
def home():
    return render_template('index.html')

def generate_policy_content(form_data):
    """Generate policy content using CrimsonAI hub"""
    try:
        prompt = f"""You are a cybersecurity policy generator trained to write business-grade IT and security policies aligned with industry frameworks such as NIST 800-53 Rev. 5, SOC 2, SEC Cyber Risk Guidance, and CIS Controls. Your policies must follow this format:
- A clear section heading for each topic
- A short introductory paragraph explaining the purpose or intent of that section
- A mix of both: lists of actionable bullet points or paragraphs that reflect standards and expected behaviors

The language should be professional, clear, and non-technical. Avoid legal jargon or redundant statements.

IMPORTANT FORMATTING REQUIREMENTS:
- Do NOT include any conversational responses like "Certainly", "Here is", etc.
- Use ONLY H1 and H2 headings, no bold text (**text**)  
- Do NOT use em dashes (—), use regular hyphens (-) instead
- Start directly with the policy content, no introductory text
- Use simple bullet points with hyphens (-) not special characters
- Keep formatting clean and professional

Generate a comprehensive {form_data['policy_type']} for the following organization:

Client: {form_data['client_name']}
Industry: {form_data['industry']}
Company Size: {form_data['company_size']}
Technology Stack: {form_data['tech_stack']}
Compliance Requirements: {form_data['compliance_requirements'] or 'General best practices'}
Managed Service Provider (MSP): Crimson IT
Managed Security Service Provider (MSSP): Crimson IT
Additional Requirements: {form_data['additional_requirements']}

The policy should be specifically tailored to this organization's context and include:
1. Purpose and scope
2. Roles and responsibilities
3. Policy statements with specific controls
4. Implementation guidelines
5. Compliance and monitoring requirements
6. References to applicable frameworks

Make sure to reference Crimson IT as the designated MSP and MSSP throughout the policy where appropriate."""

        configure_openai()
        
        # Try your exact Azure deployment names first
        model_names = ["gpt-5-chat", "gpt-4o"]
        
        response = None
        last_error = None
        
        for model_name in model_names:
            try:
                response = openai.ChatCompletion.create(
                    engine=model_name,
                    messages=[
                        {"role": "system", "content": "You are an expert cybersecurity policy writer with deep knowledge of NIST frameworks, SOC 2, ISO 27001, CMMC, and industry best practices."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=4000,
                    temperature=0.7
                )
                break  # Success, exit the loop
            except Exception as model_error:
                last_error = f"Model {model_name}: {str(model_error)}"
                continue  # Try next model
        
        if response is None:
            raise Exception(f"All models failed. Last error: {last_error}")
        
        content = response.choices[0].message.content
        
        # Clean up the response - remove conversational elements and formatting issues
        # Remove common AI prefixes
        prefixes_to_remove = [
            "Certainly!", "Certainly,", "Certainly.", "Sure!", "Sure,", "Of course!", "Of course,",
            "Here is", "Here's", "Below is", "I'll", "I can", "Let me", "I'd be happy to",
            "Here you go", "Absolutely!", "Absolutely,", "No problem!", "No problem,",
            "I have aligned", "I have created", "This policy", "The following"
        ]
        
        for prefix in prefixes_to_remove:
            if content.strip().startswith(prefix):
                content = content.strip()[len(prefix):].strip()
                break
        
        # Remove em dashes and replace with hyphens
        content = content.replace('—', '-')
        content = content.replace('–', '-')
        
        # Remove bold formatting markers
        content = content.replace('**', '')
        
        # Clean up bullet points - standardize to hyphens
        content = content.replace('•', '-')
        content = content.replace('◦', '-')
        content = content.replace('▪', '-')
        
        # Remove common AI suffixes and additional offers
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            # Skip lines that contain AI conversation elements
            if any(phrase in line.lower() for phrase in [
                "if you'd like", "would you like", "let me know if", "feel free to",
                "i can also", "i'd be happy to", "please let me know",
                "if you need", "would you prefer", "shall i", "do you want",
                "comprehensive information security policy", "tailored to the organization",
                "i have aligned", "aligned it with", "below is a", "here is a"
            ]):
                continue
            # Skip lines that are just formatting artifacts
            if line in ['--', '____________________________________________________________', '•', '◦', '▪']:
                continue
            # Skip duplicate title lines
            if line.lower().startswith('abilityfirst') and 'information security policy' in line.lower():
                continue
            if line:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
        
    except Exception as e:
        return f"Error generating policy content: {str(e)}"

def create_word_document(form_data, policy_content):
    """Create a professional Word document"""
    try:
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # Add header with Crimson IT branding
        header_para = doc.add_paragraph()
        header_run = header_para.add_run('CRIMSON IT')
        header_run.bold = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add main title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(form_data['client_name'].upper())
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(form_data['policy_type'])
        subtitle_run.bold = True
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation info
        info_para = doc.add_paragraph()
        info_para.add_run(f'Document Generated: {datetime.now().strftime("%B %d, %Y")}').italic = True
        info_para.add_run('\nManaged by: Crimson IT (MSP/MSSP)').italic = True
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        doc.add_paragraph('_' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Add policy content
        lines = policy_content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    doc.add_paragraph()
                current_paragraph = None
                continue
                
            if line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
                current_paragraph = None
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
                current_paragraph = None
            elif line.endswith(':') or (line.isupper() and len(line) > 10):
                doc.add_heading(line.rstrip(':'), level=2)
                current_paragraph = None
            elif line.startswith('-') or line.startswith('•') or line.startswith('* '):
                bullet_text = line[1:].strip() if line.startswith('-') else line[2:].strip()
                doc.add_paragraph(bullet_text, style='List Bullet')
                current_paragraph = None
            elif line.startswith(tuple('123456789')) and ('. ' in line[:5] or ') ' in line[:5]):
                doc.add_paragraph(line, style='List Number')
                current_paragraph = None
            else:
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                if current_paragraph.text:
                    current_paragraph.add_run(' ')
                current_paragraph.add_run(line)
        
        # Add footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.add_run('This document was generated by Crimson IT\'s AI-powered policy generator.').italic = True
        footer_para.add_run('\nFor questions or updates, contact your Crimson IT representative.').italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc
        
    except Exception as e:
        raise Exception(f"Error creating Word document: {str(e)}")

@app.route('/generate_policy', methods=['POST'])
def generate_policy():
    try:
        # Get form data
        form_data = {
            'client_name': request.form.get('client_name', '').strip(),
            'industry': request.form.get('industry', '').strip(),
            'company_size': request.form.get('company_size', '').strip(),
            'policy_type': request.form.get('policy_type', '').strip(),
            'compliance_requirements': request.form.get('compliance_requirements', '').strip(),
            'additional_requirements': request.form.get('additional_requirements', '').strip()
        }
        
        # Collect technology stack from individual dropdown fields
        tech_components = []
        tech_fields = [
            'platform_choice', 'mdr_solution', 'email_security', 'siem_solution',
            'pam_solution', 'disk_encryption', 'mdm_computers', 'mdm_mobile', 'security_training',
            'phishing_tests', 'vulnerability_scans', 'dark_web_monitoring', 'mfa_solution', 
            'password_manager', 'intrusion_detection', 'additional_tech'
        ]
        
        for field in tech_fields:
            value = request.form.get(field, '').strip()
            if value and value != 'None':
                tech_components.append(value)
        
        # Combine all technology selections into tech_stack
        form_data['tech_stack'] = ', '.join(tech_components) if tech_components else 'Standard business technology stack'
        
        # Validate required fields
        required_fields = ['client_name', 'industry', 'company_size', 'policy_type', 'tech_stack']
        missing_fields = [field for field in required_fields if not form_data[field]]
        
        if missing_fields:
            return jsonify({'error': f'Missing required fields: {", ".join(missing_fields)}'}), 400
        
        # Generate policy content using CrimsonAI
        policy_content = generate_policy_content(form_data)
        
        if policy_content.startswith("Error generating policy content:"):
            return jsonify({'error': policy_content}), 500
        
        # Create Word document
        doc = create_word_document(form_data, policy_content)
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            
            # Generate safe filename
            safe_client_name = "".join(c for c in form_data['client_name'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_policy_type = "".join(c for c in form_data['policy_type'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
            timestamp = datetime.now().strftime('%Y%m%d_%H%M')
            filename = f"{safe_client_name}_{safe_policy_type}_{timestamp}.docx"
            
            return send_file(
                tmp_file.name,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
    except Exception as e:
        return jsonify({'error': f'An error occurred: {str(e)}'}), 500

@app.route('/health')
def health():
    try:
        # Check if environment variables are available
        api_key = os.getenv('AZURE_OPENAI_API_KEY')
        endpoint = os.getenv('AZURE_OPENAI_ENDPOINT')
        
        health_info = {
            "status": "healthy", 
            "version": "2.0.2",
            "openai_configured": bool(api_key and endpoint),
            "env_check": {
                "api_key_present": bool(api_key),
                "endpoint_present": bool(endpoint)
            }
        }
        
        # Test OpenAI connection if configured
        if api_key and endpoint:
            try:
                configure_openai()
                # Try a simple test with each model
                test_results = {}
                model_names = ["gpt-5-chat", "gpt-4o"]
                
                for model in model_names:
                    try:
                        response = openai.ChatCompletion.create(
                            engine=model,
                            messages=[{"role": "user", "content": "Hello"}],
                            max_tokens=5
                        )
                        test_results[model] = "available"
                    except Exception as e:
                        test_results[model] = f"error: {str(e)}"
                
                health_info["model_availability"] = test_results
                
            except Exception as e:
                health_info["openai_test_error"] = str(e)
        
        return health_info
    except Exception as e:
        return {"status": "error", "error": str(e)}, 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port, debug=False)